"""DOCX text extraction via python-docx (#59 / Decision #007).

Soft-fail only — never raise into the index worker for bad files.
Legacy ``.doc`` is not supported (not registered).
"""

from __future__ import annotations

import time
from pathlib import Path

from indexer.extract import ExtractResult

MAX_DOCX_BYTES = 50 * 1024 * 1024  # 50 MiB
MAX_DOCX_EXTRACT_SECONDS = 30.0


def _python_docx_version() -> str | None:
    try:
        import docx

        return getattr(docx, "__version__", None)
    except Exception:  # noqa: BLE001
        return None


def _paragraph_texts(doc) -> list[str]:
    lines: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append("\t".join(cells))
    return lines


def extract_docx(
    path: Path | str,
    *,
    max_seconds: float | None = None,
) -> ExtractResult:
    """
    Extract paragraph + table cell text from a ``.docx`` into one FTS segment.

    Caps: skip files larger than MAX_DOCX_BYTES. Soft-fails on missing,
    corrupt, or timed-out extracts.
    """
    resolved = Path(path)
    version = _python_docx_version()
    result = ExtractResult(parser="python-docx", parser_version=version, page_count=1)
    budget = MAX_DOCX_EXTRACT_SECONDS if max_seconds is None else float(max_seconds)

    try:
        size = resolved.stat().st_size
    except OSError as exc:
        result.status = "error"
        result.warnings.append(f"stat failed: {exc}")
        result.page_count = 0
        return result

    if size > MAX_DOCX_BYTES:
        result.status = "error"
        result.warnings.append(f"file exceeds {MAX_DOCX_BYTES} bytes; skipped extract")
        result.page_count = 0
        return result

    try:
        from docx import Document
    except ImportError as exc:
        result.status = "error"
        result.warnings.append(f"python-docx not installed: {exc}")
        result.page_count = 0
        return result

    started = time.monotonic()
    try:
        doc = Document(resolved)
        if time.monotonic() - started >= budget:
            result.status = "error"
            result.warnings.append(f"extract exceeded {budget:g}s during open; stopped")
            result.page_count = 0
            return result

        lines = _paragraph_texts(doc)
        if time.monotonic() - started >= budget:
            result.status = "error"
            result.warnings.append(f"extract exceeded {budget:g}s; stopped early")
            # Keep whatever we got if any.
            if lines:
                text = "\n".join(lines)
                result.pages = [(1, text)]
                result.page_count = 1
            else:
                result.page_count = 0
            return result

        text = "\n".join(lines)
        if not text.strip():
            result.status = "empty"
            result.warnings.append("no extractable text")
            result.pages = []
            result.page_count = 0
            return result

        result.pages = [(1, text)]
        result.page_count = 1
        result.status = "ok"
        return result
    except Exception as exc:  # noqa: BLE001
        result.status = "error"
        result.warnings.append(f"extract failed: {exc}")
        result.pages = []
        result.page_count = 0
        return result
