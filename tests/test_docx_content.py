"""DOCX extract + content sync (#59)."""

from pathlib import Path

import pytest

from db import connect, init_db
from indexer import ensure_root, replace_root_files
from indexer.docx_extract import extract_docx
from indexer.extract import extract_for_path
from indexer.search import search_filenames


@pytest.fixture()
def temp_db(tmp_path, monkeypatch):
    db_path = tmp_path / "test-index.db"
    monkeypatch.setenv("AIDESKTOP_DB", str(db_path))
    init_db(db_path)
    return db_path


def _write_docx(path: Path, paragraphs: list[str], *, table_rows: list[list[str]] | None = None) -> Path:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell
    doc.save(path)
    return path


def test_extract_docx_paragraphs(tmp_path: Path):
    path = _write_docx(tmp_path / "note.docx", ["hello gyarados world"])
    result = extract_docx(path)
    assert result.status == "ok"
    assert result.parser == "python-docx"
    assert result.page_count == 1
    assert "gyarados" in result.pages[0][1].lower()


def test_extract_docx_tables(tmp_path: Path):
    path = _write_docx(
        tmp_path / "table.docx",
        ["intro"],
        table_rows=[["alpha", "lapras"], ["beta", "two"]],
    )
    result = extract_docx(path)
    assert result.status == "ok"
    assert "lapras" in result.pages[0][1].lower()


def test_extract_docx_empty(tmp_path: Path):
    path = _write_docx(tmp_path / "empty.docx", [])
    result = extract_docx(path)
    assert result.status == "empty"
    assert result.pages == []


def test_extract_docx_corrupt_soft_fails(tmp_path: Path):
    path = tmp_path / "bad.docx"
    path.write_bytes(b"not a real docx archive")
    result = extract_docx(path)
    assert result.status == "error"
    assert result.warnings


def test_extract_for_path_dispatches_docx(tmp_path: Path):
    path = _write_docx(tmp_path / "via.docx", ["registry docx snorlax"])
    result = extract_for_path(path)
    assert result.status == "ok"
    assert result.parser == "python-docx"
    assert "snorlax" in result.pages[0][1].lower()


def test_docx_sync_and_search(temp_db, tmp_path: Path):
    folder = tmp_path / "docs"
    path = _write_docx(
        folder / "charizard.docx",
        ["Title is Charizard; body mentions unique token mewtwo."],
    )
    root_id = ensure_root(folder)
    replace_root_files(root_id, [path])

    with connect() as conn:
        file_id = int(
            conn.execute(
                "SELECT id FROM files WHERE path = ?",
                (str(path.resolve()),),
            ).fetchone()["id"]
        )
        content = conn.execute(
            "SELECT status, parser, page_count FROM file_content WHERE file_id = ?",
            (file_id,),
        ).fetchone()
        pages = conn.execute(
            "SELECT page, text FROM file_pages_fts WHERE file_id = ?",
            (file_id,),
        ).fetchall()

    assert content["status"] == "ok"
    assert content["parser"] == "python-docx"
    assert int(content["page_count"]) == 1
    assert len(pages) == 1
    assert pages[0]["page"] == 1
    assert "mewtwo" in pages[0]["text"].lower()

    hits = search_filenames("mewtwo")
    assert len(hits) == 1
    assert hits[0]["name"] == "charizard.docx"
    assert hits[0]["match"] in ("content", "both")
    assert hits[0]["page"] == 1
    assert hits[0]["extension"] == "docx"
